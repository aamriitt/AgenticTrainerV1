"""
Agent 1: Data Ingestion Agent
=============================

Responsibility
--------------
Accept raw enterprise files (PDF, DOCX, TXT, FAQ, video) and turn each
into a `RawDocument`: plain text plus enough structural metadata
(page numbers, timestamps) that later agents can build precise
citations.

This agent does NOT clean or chunk text — that's Agent 2 and Agent 3.
Its only job is faithful extraction.

Design notes
------------
- Each file type gets its own small, testable function. `ingest_file`
  dispatches by extension so callers don't need to know the details.
- Video ingestion requires Whisper + ffmpeg locally. If Whisper isn't
  installed/available, `ingest_video` raises a clear, actionable
  error rather than failing silently.
- OCR (Tesseract) is used as a fallback for PDF pages that contain no
  extractable text layer (i.e. scanned pages/images).
"""

from __future__ import annotations

import io
from pathlib import Path

from app.utils.logger import get_logger
from app.utils.schemas import RawDocument, SourceType

logger = get_logger("agents.ingestion")

SUPPORTED_EXTENSIONS = {
    ".pdf": SourceType.PDF,
    ".docx": SourceType.DOCX,
    ".txt": SourceType.TXT,
    ".mp4": SourceType.VIDEO,
    ".mov": SourceType.VIDEO,
    ".mkv": SourceType.VIDEO,
    ".wav": SourceType.VIDEO,
    ".mp3": SourceType.VIDEO,
}


class IngestionError(Exception):
    """Raised when a file cannot be ingested."""


def _source_id_from_path(path: Path) -> str:
    return path.stem


def ingest_pdf(file_path: str, ocr_fallback: bool = True) -> RawDocument:
    """
    Extract text from a PDF, page by page.

    Uses PyMuPDF for the text layer. If a page has little/no
    extractable text (common with scanned documents), falls back to
    OCR via pytesseract on a rasterized image of that page.
    """
    try:
        import fitz  # PyMuPDF
    except ImportError as e:
        raise IngestionError(
            "PyMuPDF is not installed. Run: pip install PyMuPDF"
        ) from e

    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File not found: {file_path}")

    segments: dict[str, str] = {}
    full_text_parts: list[str] = []

    doc = fitz.open(file_path)
    try:
        for page_index in range(len(doc)):
            page = doc.load_page(page_index)
            page_num = page_index + 1
            text = page.get_text("text").strip()

            if not text and ocr_fallback:
                text = _ocr_page(page)
                if text:
                    logger.info(f"{path.name} page {page_num}: used OCR fallback")

            if text:
                segments[f"Page {page_num}"] = text
                full_text_parts.append(text)
    finally:
        doc.close()

    if not full_text_parts:
        logger.warning(f"No extractable text found in {path.name}")

    return RawDocument(
        source_id=_source_id_from_path(path),
        source_type=SourceType.PDF,
        file_path=str(path),
        text="\n\n".join(full_text_parts),
        segments=segments,
        extra_metadata={"num_pages": len(segments)},
    )


def _ocr_page(page) -> str:
    """Rasterize a PyMuPDF page and run Tesseract OCR on it."""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        logger.warning("pytesseract/Pillow not installed; skipping OCR fallback")
        return ""

    try:
        pix = page.get_pixmap(dpi=200)
        img = Image.open(io.BytesIO(pix.tobytes("png")))
        return pytesseract.image_to_string(img).strip()
    except Exception as e:  # OCR failures shouldn't crash the whole pipeline
        logger.warning(f"OCR failed on a page: {e}")
        return ""


def ingest_docx(file_path: str) -> RawDocument:
    """
    Extract text from a DOCX file, preserving paragraph/heading
    structure well enough to build "Section X" style citations.
    """
    try:
        import docx  # python-docx
    except ImportError as e:
        raise IngestionError(
            "python-docx is not installed. Run: pip install python-docx"
        ) from e

    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File not found: {file_path}")

    document = docx.Document(file_path)

    segments: dict[str, str] = {}
    full_text_parts: list[str] = []
    current_section = "Introduction"
    section_buffer: list[str] = []
    section_index = 0

    def flush_section():
        nonlocal section_buffer, section_index
        if section_buffer:
            section_index += 1
            key = f"Section {section_index}: {current_section}"
            joined = "\n".join(section_buffer)
            segments[key] = joined
            full_text_parts.append(joined)
            section_buffer = []

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        style = (para.style.name or "").lower()
        if "heading" in style:
            flush_section()
            current_section = text
        else:
            section_buffer.append(text)

    flush_section()

    # Also pull table content, since SOPs/manuals often use tables.
    table_texts = []
    for t_index, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            rows.append(" | ".join(cells))
        if rows:
            table_text = "\n".join(rows)
            segments[f"Table {t_index}"] = table_text
            table_texts.append(table_text)

    full_text_parts.extend(table_texts)

    return RawDocument(
        source_id=_source_id_from_path(path),
        source_type=SourceType.DOCX,
        file_path=str(path),
        text="\n\n".join(full_text_parts),
        segments=segments,
        extra_metadata={"num_sections": section_index, "num_tables": len(table_texts)},
    )


def _split_faq_pairs(text: str) -> dict[str, str]:
    """
    Split an FAQ file into individual Q/A segments so each pair stays
    intact through chunking. Expects a loose "Q: ... A: ..." format
    (blank-line separated). Falls back to an empty dict (caller uses
    a single "Full Document" segment) if no Q:/A: markers are found.
    """
    import re

    blocks = re.split(r"\n\s*\n", text)
    qa_blocks = [b.strip() for b in blocks if b.strip()]

    if not qa_blocks or not any(re.search(r"^Q\s*:", b, re.MULTILINE) for b in qa_blocks):
        return {}

    segments: dict[str, str] = {}
    for i, block in enumerate(qa_blocks, start=1):
        match = re.search(r"^Q\s*:\s*(.+)$", block, re.MULTILINE)
        label = f"FAQ #{i}"
        if match:
            question_preview = match.group(1).strip()[:60]
            label = f"FAQ #{i}: {question_preview}"
        segments[label] = block

    return segments


def ingest_txt(file_path: str) -> RawDocument:
    """
    Plain text / FAQ files.

    FAQ files (anything under a `faq/` folder) are split into one
    segment per question/answer pair so a pair is never split apart
    during chunking. Other TXT files are treated as a single segment.
    """
    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File not found: {file_path}")

    text = path.read_text(encoding="utf-8", errors="replace").strip()
    source_type = SourceType.FAQ if "faq" in path.parts else SourceType.TXT

    segments: dict[str, str] = {}
    if text:
        if source_type == SourceType.FAQ:
            segments = _split_faq_pairs(text)
        if not segments:
            segments = {"Full Document": text}

    return RawDocument(
        source_id=_source_id_from_path(path),
        source_type=source_type,
        file_path=str(path),
        text=text,
        segments=segments,
        extra_metadata={"num_qa_pairs": len(segments)} if source_type == SourceType.FAQ else {},
    )


def ingest_video(file_path: str, model_size: str = "base") -> RawDocument:
    """
    Transcribe a video/audio file with Whisper and segment the
    transcript by timestamp range so citations can point to
    "00:14:32" style locations.

    NOTE: Requires the `openai-whisper` package and ffmpeg installed
    on the host machine, plus enough compute to run inference
    (CPU works but is slow; GPU strongly recommended for anything
    beyond short clips).
    """
    try:
        import whisper
    except ImportError as e:
        raise IngestionError(
            "openai-whisper is not installed. Run: pip install openai-whisper "
            "(and ensure ffmpeg is installed on your system)."
        ) from e

    path = Path(file_path)
    if not path.exists():
        raise IngestionError(f"File not found: {file_path}")

    logger.info(f"Loading Whisper model '{model_size}' (first run downloads weights)...")
    model = whisper.load_model(model_size)

    result = model.transcribe(str(path))

    segments: dict[str, str] = {}
    full_text_parts: list[str] = []

    for seg in result.get("segments", []):
        start = _format_timestamp(seg["start"])
        end = _format_timestamp(seg["end"])
        seg_text = seg["text"].strip()
        if seg_text:
            segments[f"{start}-{end}"] = seg_text
            full_text_parts.append(seg_text)

    return RawDocument(
        source_id=_source_id_from_path(path),
        source_type=SourceType.VIDEO,
        file_path=str(path),
        text=" ".join(full_text_parts),
        segments=segments,
        extra_metadata={"language": result.get("language"), "num_segments": len(segments)},
    )


def _format_timestamp(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def ingest_file(file_path: str, whisper_model_size: str = "base") -> RawDocument:
    """
    Dispatch to the correct ingestion function based on file
    extension. This is the single entry point the rest of the
    pipeline (or an upload API endpoint) should call.
    """
    path = Path(file_path)
    ext = path.suffix.lower()

    if ext not in SUPPORTED_EXTENSIONS:
        raise IngestionError(
            f"Unsupported file type '{ext}'. Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )

    logger.info(f"Ingesting {path.name} ({ext})")

    if ext == ".pdf":
        return ingest_pdf(str(path))
    if ext == ".docx":
        return ingest_docx(str(path))
    if ext == ".txt":
        return ingest_txt(str(path))
    if ext in {".mp4", ".mov", ".mkv", ".wav", ".mp3"}:
        return ingest_video(str(path), model_size=whisper_model_size)

    # Should be unreachable given the check above.
    raise IngestionError(f"No handler implemented for '{ext}'")
